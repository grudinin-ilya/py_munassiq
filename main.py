import os
import shutil
import subprocess
import sys
import design
import stdDatabase_BackEnd

from PyQt5 import QtWidgets, QtGui

from PyQt5.QtWidgets import QTableWidgetItem, QMessageBox, QFrame, QLabel

from docx import Document
# from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.shared import Inches, Pt
from firebase import firebase

firebase = firebase.FirebaseApplication('', None)

# вставмть в design перед функцией
# ===============
# import datetime
# from PyQt5.QtCore import QRegExp
# from PyQt5.QtGui import QRegExpValidator
# self.dateEdit.setDate(datetime.date.today())
#
# number = QRegExp("[0-9]+")
# validator_number = QRegExpValidator(number)
# address = QRegExp("[0-9/]+")
# validator_address = QRegExpValidator(address)
# letter = QRegExp("[\u0621-\u064A ]+")
# validator_letter = QRegExpValidator(letter)
#
# self.lineEdit_FName.setValidator(validator_letter)
# self.lineEdit_Nationality.setValidator(validator_letter)
# self.lineEdit_StudID.setValidator(validator_number)
# self.lineEdit_Faculty.setValidator(validator_letter)
# self.lineEdit_Address.setValidator(validator_address)
# self.lineEdit_Mobile.setValidator(validator_number)
# self.lineEdit_Whats.setValidator(validator_number)
# self.lineEdit_Disability.setValidator(validator_letter)
# self.lineEdit_unit.setValidator(validator_number)
#     ====================

basmala = 'بسم الله الحمن الرحيم' + '\n'
iu = 'الجامعة الإسلامية بالمدينة المنورة' + '\n'
udis = 'وحدة ذوي الإعاقة' + '\n'
head = 'سعادة المشرف على وحدة ذوي الإعاقة' + '\t\t\t' + 'وفقه الله' + '\n'
salam = 'السلام عليكم و رحمة الله و بركاته' + '\t\t\t' + 'و بعد:' + '\n'
letter = 'فأخبركم بأن الطالب'


# style = document.styles.add_style("table-rtl", WD_STYLE_TYPE.TABLE)
# style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

class VLine(QFrame):
    def __init__(self):
        super(VLine, self).__init__()
        self.setFrameShape(self.VLine | self.Sunken)


class ExampleApp(QtWidgets.QMainWindow, design.Ui_MainWindow):

    def __init__(self):

        super().__init__()
        self.setupUi(self)

        # StatusBar()

        self.lbl1 = QLabel("وحدة ذوي الإعاقة")
        self.lbl2 = QLabel("الجامعة الإسلامية بالمدينة المنورة")

        self.statusBar().reformat()
        self.statusBar().setStyleSheet('border: 0; background-color: #FFF8DC;')
        self.statusBar().setStyleSheet("QStatusBar::item {border: none;}")

        self.statusBar().addPermanentWidget(VLine())  # <---
        self.statusBar().addPermanentWidget(self.lbl1)
        self.statusBar().addPermanentWidget(VLine())  # <---
        self.statusBar().addPermanentWidget(self.lbl2)
        self.statusBar().addPermanentWidget(VLine())  # <---

        # style

        self.setStyleSheet('QMainWindow{background-color: darkgray;border: 1px solid black;}')

        # кнопки и их привязка к функциям. первая часть программы
        self.pushButton_addData.clicked.connect(self.add_data)
        self.pushButton_cleanData.clicked.connect(self.clean_data)
        self.pushButton_displayData.clicked.connect(self.display_data)
        self.tableWidget.itemClicked.connect(self.clicked_item_table)
        self.pushButton_deleteData.clicked.connect(self.delete_data)
        self.pushButton_searchData.clicked.connect(self.search_data)
        self.pushButton_updateData.clicked.connect(self.update_data)
        self.pushButton_folder_stud.clicked.connect(self.open_folder_stud)
        self.pushButton_manual_tab1.clicked.connect(self.manual_tab1)
        self.pushButton_copy.clicked.connect(self.copy)

        # вторая часть программы

        self.comboBox.currentIndexChanged.connect(self.selectionchange)
        self.pushButton_updateComboBox.clicked.connect(self.update_combo_box)
        self.pushButton_saveVisit.clicked.connect(self.save_visit)
        self.pushButton_saveReport.clicked.connect(self.save_report)
        self.pushButton_quit.clicked.connect(self.quit)
        self.pushButton_manual.clicked.connect(self.manual)
        self.pushButton_folder_request.clicked.connect(self.open_folder_request)
        self.pushButton_folder_visit.clicked.connect(self.open_folder_visit)
        self.pushButton_manual_tab2.clicked.connect(self.manual_tab2)
        self.pushButton_add_names_firebase.clicked.connect(self.add_names_firebase)
        self.pushButton_download_result_firebase.clicked.connect(self.download_result_firebase)

    def add_data(self):
        name_stud = self.lineEdit_FName.text()
        nationality = self.lineEdit_Nationality.text()
        stud_id = self.lineEdit_StudID.text()
        faculty = self.lineEdit_Faculty.text()
        address = self.lineEdit_Address.text()
        mobile = self.lineEdit_Mobile.text()
        whats = self.lineEdit_Whats.text()
        disability = self.lineEdit_Disability.text()
        if len(name_stud) != 0:
            stdDatabase_BackEnd.add_std_rec(name_stud, nationality, stud_id, faculty,
                                            address, mobile, whats, disability)
            result = stdDatabase_BackEnd.view(name_stud)
            self.show_table(result)
            self.clean_data()
            self.save_database_in_docx()

    def clean_data(self):
        self.lineEdit_FName.setText("")
        self.lineEdit_Nationality.setText("")
        self.lineEdit_StudID.setText("")
        self.lineEdit_Faculty.setText("")
        self.lineEdit_Address.setText("")
        self.lineEdit_Mobile.setText("")
        self.lineEdit_Whats.setText("")
        self.lineEdit_Disability.setText("")

    def show_table(self, result):
        self.tableWidget.setRowCount(0)
        for row_number, row_data in enumerate(result):
            self.tableWidget.insertRow(row_number)
            for column_number, data in enumerate(row_data):
                self.tableWidget.setItem(row_number, column_number, QTableWidgetItem(str(data)))

    def display_data(self):
        result = stdDatabase_BackEnd.view_data()
        self.show_table(result)

    def clicked_item_table(self, item):
        global item_table
        item_table = item.text()
        for i in stdDatabase_BackEnd.view(item_table):
            self.lineEdit_FName.setText(i[0])
            self.lineEdit_Nationality.setText(i[1])
            self.lineEdit_StudID.setText(i[2])
            self.lineEdit_Faculty.setText(i[3])
            self.lineEdit_Address.setText(i[4])
            self.lineEdit_Mobile.setText(i[5])
            self.lineEdit_Whats.setText(i[6])
            self.lineEdit_Disability.setText(i[7])

    def delete_data(self):
        if not self.tableWidget.rowCount() or (self.tableWidget.rowCount() and len(self.lineEdit_FName.text()) == 0):
            self.create_msg_box('/how_del.html', 'تنبيه', QtWidgets.QMessageBox.Ok, '/msg_error.png')
        elif self.tableWidget.rowCount() and len(self.lineEdit_FName.text()) != 0:
            content = self.create_msg_box('/del_stud.html', 'تنبيه',
                                          QtWidgets.QMessageBox.Ok | QtWidgets.QMessageBox.No,
                                          '/msg_data_cleaning.png')
            if content == QtWidgets.QMessageBox.Ok:
                stdDatabase_BackEnd.delete_rec(self.lineEdit_FName.text())
                self.clean_data()
                self.display_data()
                self.save_database_in_docx()

    def search_data(self):
        if len(self.lineEdit_Address.text()) != 0:
            result = stdDatabase_BackEnd.search_address(self.lineEdit_Address.text())
        else:
            result = stdDatabase_BackEnd.search_data(self.lineEdit_FName.text(), self.lineEdit_Nationality.text(),
                                                     self.lineEdit_StudID.text(), self.lineEdit_Faculty.text(),
                                                     self.lineEdit_Mobile.text(),
                                                     self.lineEdit_Whats.text(), self.lineEdit_Disability.text())
        self.show_table(result)

    def update_data(self):
        if len(self.lineEdit_FName.text()) == 0:
            self.create_msg_box('/how_update.html', 'تنبيه', QtWidgets.QMessageBox.Ok, '/msg_os.png')
        elif len(self.lineEdit_FName.text()) != 0:
            stdDatabase_BackEnd.delete_rec(self.lineEdit_FName.text())
            self.add_data()

    def save_database_in_docx(self):
        document = Document()
        head = basmala + iu + udis + 'بيانات الطلاب المسجلين في وحدة ذوي الإعاقة الساكنين داخل الجامعة'
        document.add_paragraph(head)
        table = document.add_table(rows=1, cols=8)

        records = stdDatabase_BackEnd.view_data()

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'الإعاقة'
        # hdr_cells[0].width = Inches(2.5)
        hdr_cells[1].text = 'رقم الواتس'
        # hdr_cells[1].width = Inches(0.5)
        hdr_cells[2].text = 'رقم الجوال'
        # hdr_cells[2].width = Inches(1.5)
        hdr_cells[3].text = 'السكن'
        # hdr_cells[3].width = Inches(2.5)
        hdr_cells[4].text = 'جهة التعليم'
        hdr_cells[5].text = 'رقم الطالب'
        hdr_cells[6].text = 'الجنسية'
        hdr_cells[7].text = 'الاسم'

        for name_stud, nationality, stud_id, faculty, address, mobile, whats, disability, in records:
            row_cells = table.add_row().cells
            row_cells[7].text = str(name_stud)
            row_cells[6].text = nationality
            row_cells[5].text = stud_id
            row_cells[4].text = faculty
            row_cells[3].text = address
            row_cells[2].text = mobile
            row_cells[1].text = whats
            row_cells[0].text = disability
        document.add_page_break()
        file_name = 'قائمة البيانات.docx'
        try:
            self.del_exist('/Stud', file_name)
        except Exception:
            pass
        document.save(file_name)

        self.move_file(os.getcwd(), '/Stud', file_name)

    def del_exist(self, folder, file_name):

        os.remove(os.path.join(os.getcwd() + folder, file_name))

    def copy(self):
        import pyperclip
        text = 'اسم الطالب: ' + self.lineEdit_FName.text() + '\n' + 'الجنسية: ' + self.lineEdit_Nationality.text() + '\n' + 'رقم الطالب: ' + self.lineEdit_StudID.text() + '\n' + 'جهة التعليم: ' + self.lineEdit_Faculty.text() + '\n' + 'السكن: ' + self.lineEdit_Address.text() + '\n' + 'الجوال: ' + self.lineEdit_Mobile.text() + '\n' + 'الواتس: ' + self.lineEdit_Whats.text() + '\n' + 'الإعاقة: ' + self.lineEdit_Disability.text()

        pyperclip.copy(text)

    # ==================функции второй части программы
    def clear_choice(self):
        self.textEdit_need.setText("")
        self.buttonGroup.setExclusive(False)
        self.radioButton_attendance.setChecked(False)
        self.radioButton_notNead.setChecked(False)
        self.buttonGroup.setExclusive(True)

    def selectionchange(self):
        self.clear_choice()

    def update_combo_box(self):
        if len(self.lineEdit_unit.text()) != 0:
            self.comboBox.clear()
            self.comboBox.addItems([""])
            self.comboBox.addItems(stdDatabase_BackEnd.get_name(self.lineEdit_unit.text())[0])
        else:
            self.create_msg_box('/no_unit.html', 'ادخل رقم الوحدة', QtWidgets.QMessageBox.Ok, '/msg_alert.png')

    def save_visit(self):
        name_stud = self.comboBox.currentText()

        if len(name_stud) != 0:
            items = stdDatabase_BackEnd.get_items(name_stud)
            if self.radioButton_attendance.isChecked() and len(self.textEdit_need.toPlainText()) == 0:
                result_visit = "غائب"
            elif self.radioButton_notNead.isChecked() and len(self.textEdit_need.toPlainText()) == 0:
                result_visit = "لا يحتاج إلى أي شيء حاليا"
            elif len(self.textEdit_need.toPlainText()) != 0 and (
                    not self.radioButton_attendance.isChecked() and not self.radioButton_notNead.isChecked()):
                result_visit = self.textEdit_need.toPlainText()
                self.save_request(name_stud, self.textEdit_need.toPlainText())
            if (self.radioButton_attendance.isChecked() or self.radioButton_notNead.isChecked()) and len(
                    self.textEdit_need.toPlainText()) != 0:
                self.create_msg_box("/two_choice.html", "للمعلومة", QtWidgets.QMessageBox.Ok, "/msg_alert.png")
                self.clear_choice()
        else:
            self.create_msg_box("/chooce_stud.html", 'تنبيه', QtWidgets.QMessageBox.Ok, "/msg_ui.png")
            self.clear_choice()
        try:
            stdDatabase_BackEnd.add_visit(self.dateEdit.text(), name_stud, items[0], items[1], result_visit)
        except Exception:
            pass

    def save_report(self):
        document = Document()
        head = basmala + iu + udis + 'تقرير الزيارة ' + '\t' + self.dateEdit.text()
        document.add_paragraph(head)
        table = document.add_table(rows=1, cols=5)
        records = stdDatabase_BackEnd.view_result_visit(self.dateEdit.text())
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'نتيجة الزيارة'
        # hdr_cells[0].width = Inches(2.5)
        hdr_cells[1].text = 'السكن'
        # hdr_cells[1].width = Inches(0.5)
        hdr_cells[2].text = 'رقم الطالب'
        # hdr_cells[2].width = Inches(1.5)
        hdr_cells[3].text = 'الاسم'
        # hdr_cells[3].width = Inches(2.5)
        hdr_cells[4].text = 'تاريخ الزيارة'
        for Date, FName, StudID, Address, ResultVisit, in records:
            row_cells = table.add_row().cells
            row_cells[4].text = str(Date)
            row_cells[3].text = FName
            row_cells[2].text = StudID
            row_cells[1].text = Address
            row_cells[0].text = ResultVisit
        document.add_page_break()
        file_name = 'تقرير الزيارة' + '.' + self.dateEdit.text() + '.docx'
        try:
            self.del_exist('/Visit', file_name)
        except Exception:
            pass
        document.save(file_name)

        self.move_file(os.getcwd(), '/Visit', file_name)

    def save_request(self, name_stud, needs):
        document = Document()
        stud_id = stdDatabase_BackEnd.get_items(name_stud)[0]
        record = basmala + iu + udis + head + salam + letter + ' ' + name_stud + ' ' + '(' + stud_id + ')' + ' ' + needs + '\n' + self.dateEdit.text()
        document.add_paragraph(record)
        document.add_page_break()
        file_name = name_stud + '.' + self.dateEdit.text() + '.docx'
        try:
            self.del_exist('/Request', file_name)
        except Exception:
            pass
        document.save(file_name)

        self.move_file(os.getcwd(), '/Request', file_name)

    def open_dir(self, new_folder):
        path = os.getcwd() + new_folder
        os.path.realpath(path)
        # for linux
        return subprocess.call(['xdg-open', path])
        # for windows
        # return os.startfile(path)

    def open_folder_stud(self):
        self.open_dir('/Stud')

    def open_folder_request(self):
        self.open_dir('/Request')

    def open_folder_visit(self):
        self.open_dir('/Visit')

    def move_file(self, old_path, new_folder, file_name):
        new_path = old_path + new_folder
        shutil.move(file_name, new_path)
        # content = self.create_msg_box('/create_docx.html', file_name,
        #                               QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No,
        #                               '/msg_open_folder.png')
        # if content == QtWidgets.QMessageBox.Yes:
        #     self.open_dir(new_folder)
        self.clear_choice()

    def quit(self):
        content = self.create_msg_box('/exit.html', 'الخروج', QtWidgets.QMessageBox.Ok | QtWidgets.QMessageBox.Cancel,
                                      '/msg_exit.png')
        if content == QtWidgets.QMessageBox.Ok:
            QtWidgets.qApp.exit()

    def manual(self):
        self.create_msg_box('/manual.html', 'دليل استعمال البرنامج', QtWidgets.QMessageBox.Ok, False)

    def manual_tab1(self):
        self.create_msg_box('/manual_tab1.html', 'دليل تبويب التسجيل و البيانات', QtWidgets.QMessageBox.Ok,
                            '/msg_manual.png')

    def manual_tab2(self):
        self.create_msg_box('/manual_tab2.html', 'دليل تبويب الزيارات و التقارير', QtWidgets.QMessageBox.Ok,
                            '/msg_manual.png')

    def create_msg_box(self, html, title, button, image):
        data_dir = os.getcwd() + '/msgs'
        with open(data_dir + html, 'r', encoding="utf-8") as f:
            content = f.read()
        msg = QtWidgets.QMessageBox(self)
        data_dir = os.getcwd() + '/imgs'
        if image:
            msg.setIconPixmap(QtGui.QIcon(data_dir + image).pixmap(65))
        else:
            pass
        msg.setWindowTitle(title)
        msg.setText(content)
        msg.setStandardButtons(button)
        message = msg.exec_()
        return message

    def add_names_firebase(self):
        if len(self.lineEdit_unit.text()) != 0:
            stud_names = stdDatabase_BackEnd.get_name(self.lineEdit_unit.text())[0]
            firebase.put('/Units/', self.lineEdit_unit.text(), stud_names)

            names = stdDatabase_BackEnd.get_name(self.lineEdit_unit.text())
            name = names[0]
            address = names[1]
            result = []
            for n, a in zip(name, address):
                item = {'name': n, 'address': a}
                result.append(item)
            firebase.put('/Lists/', self.lineEdit_unit.text(), result)
        else:
            self.create_msg_box('/no_unit.html', 'ادخل رقم الوحدة', QtWidgets.QMessageBox.Ok, '/msg_alert.png')

    def download_result_firebase(self):
        if 'Results' in firebase.get('', ''):
            result = firebase.get('/Results/', '')
            for x in result.values():
                for i in x.items():
                    items = stdDatabase_BackEnd.get_items(i[0])
                    stdDatabase_BackEnd.add_visit(self.dateEdit.text(), i[0], items[0], items[1], i[1])
                    if len(i[1]) > 25:
                        self.save_request(i[0], i[1])
            firebase.delete('', '/Results/')
        else:
            self.create_msg_box('/not_found.html', 'استعمل التطبيق', QtWidgets.QMessageBox.Ok, '/msg_not_found.png')


def main():
    app = QtWidgets.QApplication(sys.argv)
    window = ExampleApp()
    window.show()
    app.exec_()


if __name__ == '__main__':
    main()

